import openpyxl
import docx

wb = openpyxl.load_workbook("./src/Momotaro.xlsx")
ws = wb.active


def pullText():
    engText = [ws[f"B{col + 1}"].value for col in range(len(ws["B"]))]
    jpText = [ws[f"C{col + 1}"].value for col in range(len(ws["C"]))]

    return engText, jpText


def createBook(text):
    # print(text)
    book = docx.Document()
    i = 0

    for i in range(0, len(text[0]), 2):
        book.add_paragraph(f"{ text[0][i]}\n{text[1][i]} \n")
        if i != len(text[0]) - 1:
            book.add_paragraph(f"{ text[0][i + 1]}\n{text[1][i + 1]} \n")
        book.add_page_break()

    book.save("./output/book.docx")


createBook(pullText())
